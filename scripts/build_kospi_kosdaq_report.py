#!/usr/bin/env python3
"""코스피·코스닥 괴리 분석 리포트를 Word(.docx)와 Excel(.xlsx)로 생성.

출력:
  research/코스피_코스닥_괴리분석_2026.docx
  research/코스피_코스닥_데이터_2026.xlsx
"""
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(BASE, "research")
os.makedirs(OUT_DIR, exist_ok=True)

NAVY = RGBColor(0x1F, 0x35, 0x64)
GRAY = RGBColor(0x55, 0x55, 0x55)

# ----------------------------------------------------------------------------
# 공통 데이터 (Word 표 + Excel 공유)
# ----------------------------------------------------------------------------
INDEX_PATH = [
    ["2025-07-01", "3,089.65", "3,000 첫 돌파(+29% YTD)", "783.67", "+14% YTD, 800 사수난항", "괴리 가시화"],
    ["2025-11-09", "6월比 +56%", "삼성·SK하이닉스 견인", "6월比 +26%", "반도체 붐 소외", "격차 30%p"],
    ["2025-12-30", "4,214.17", "연 +76% (세계 1위)", "925.47", "연 +37%", "43년래 3번째 폭등장"],
    ["2026-01-23", "—", "—", "993.43", "4년 만에 990 상회", "코스닥 반등 시작"],
    ["2026-05-22", "7,847.71", "조정", "1,161.13", "+4.99% 급등", "국민성장펀드 기대"],
    ["2026-05-29", "8,476.15", "사상 최고(전년比 ~+211%)", "1,074.8", "-2.68%", "—"],
]

DIVERGENCE = [
    ["2025 연간 수익률", "+76%", "+37%", "약 39%p"],
    ["2026 YTD(5월말)", "+86.7%", "~+24%", "약 63%p"],
    ["시가총액", "3,000조원($2T) 첫 돌파", "~400조원(5년 정체)", "—"],
    ["랠리 폭(breadth)", "지수 +24%여도 상장사 82% 하락", "—", "소수 대형주 쏠림"],
]

FOREIGN_OWN = [
    ["2020(불장 전)", "~38%", "이후 2021 ~31%로 하락"],
    ["2022", "~30%", "GFC 이후 최저"],
    ["2024 말", "32.2%", "—"],
    ["2025-07", "32.7%", "42개월 최저권"],
    ["2025 말", "36.28%", "—"],
    ["2026-05-27", "39.45%", "사상 최고(역설: 순매도에도 지분율↑)"],
]

FLOWS = [
    ["2020", "+38.94조(8월말 누적)", "-24.40조", "-13.80조", "동학개미 탄생, 외인지분 38→31%"],
    ["2021~2022", "순매수(흡수)", "구조적 순매도", "순매도", "외인지분 ~30%"],
    ["2023", "혼조", "저점 매수 복귀", "혼조", "코스닥 2차전지 버블"],
    ["2024 상반기", "—", "+22.9조(반기 사상최대)", "—", "밸류업 효과"],
    ["2024 하반기", "—", "대량 매도", "—", "반도체 고점우려+계엄쇼크, 코스닥 -21%"],
    ["2025", "코스닥 +6.2조", "~3월 -40조→5월 +5조", "코스닥 -1.4조", "HBM 모멘텀, 코스닥 외인 -1.8조"],
    ["2026 YTD", "+44.31조(5/7~22)", "코스피 -91.13조 / 코스닥 +2.84조", "—", "동학개미 redux, 예수금 ~119조"],
]

DRIVERS = [
    ["1. 반도체/AI 집중", "삼성+SK하이닉스 코스피 50%+, HBM 슈퍼사이클", "코스피 ▲▲▲ / 코스닥 ▽"],
    ["2. 밸류업+상법개정", "저PBR 대형 가치주 재평가, 밸류업지수 코스피67:코스닥33", "코스피 ▲▲ / 코스닥 ▽"],
    ["3. 금리 인하(2.5%)", "자산가격↑, 소형주 실적구제 효과 약함", "코스피 ▲ / 코스닥 ▽"],
    ["4. MSCI 선진국 기대", "패시브 자금은 대형주 매수, 2026.6 리뷰(긍정 ~60%)", "코스피 ▲ / 코스닥 —"],
    ["5. 코스닥 고유약점", "2차전지·바이오 버블붕괴, 상반기 47% 적자기업", "코스닥 ▽▽"],
]

FORECAST = [
    ["코스피", "KB증권", "10,500", "AI이익>지수, 메모리 전략자산"],
    ["코스피", "삼성자산운용", "8,000", "이익체력 퀀텀점프"],
    ["코스피", "Morgan Stanley", "6,500 (bull 7,500/bear 5,000)", "AI메모리 슈퍼사이클"],
    ["코스피", "Macquarie", "6,000", "메모리 사상최악 공급부족"],
    ["코스피", "JP Morgan", "5,000", "12개월 목표"],
    ["코스피", "Kiwoom(보수)", "4,500 (3,900~5,200)", "반도체 의존·관세 리스크"],
    ["코스닥", "유안타 등", "base 1,600 / bull 1,900", "정책+순환매"],
    ["코스닥", "삼성자산운용", "추가 +10%↑", "AI반도체 순환매"],
]

SOURCES = [
    ["코스피 3,000 돌파/괴리", "Korea Times", "2025-07-01"],
    ["2025 결산 +76%/+37%", "Korea Herald", "2025-12-30"],
    ["코스닥 정체·수급분석", "Korea Herald/The Investor", "2025-11-09"],
    ["국민성장펀드 코스닥 급등", "Korea Times", "2026-05-22"],
    ["5월 외인 코스피 사상최대 순매도", "Korea Herald/Seoul Economic Daily", "2026-05-31"],
    ["82% 하락(좁은 랠리)", "Seoul Economic Daily", "2026-05-29"],
    ["2020 동학개미 수급", "더스쿠프", "2020-09-15"],
    ["2024 상반기 외인 +22.9조", "KED Global", "2024-07-08"],
    ["외인지분율 추이", "ASIFMA / The Investor", "2023-03 / 2026-05"],
    ["밸류업 프로그램", "FSC / KRX", "2024"],
    ["상법 개정", "KED Global / Korea Herald", "2025-07·08"],
    ["BOK 금리·메모리 사이클", "CNBC", "2024-10~2026-05"],
    ["2026 코스피 전망(엇갈림)", "Korea Times / Investing.com / Bloomingbit", "2025-11~2026-05"],
    ["MSCI 편입", "KED Global", "2025-06-25"],
]

# ----------------------------------------------------------------------------
# WORD
# ----------------------------------------------------------------------------
def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i]._tc.get_or_add_tcPr()
        from docx.oxml.ns import qn
        shd = hdr[i]._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {qn("w:fill"): "1F3564"})
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    return t


def heading(doc, text, size=14, color=NAVY, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    return p


def build_word():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("코스피·코스닥 괴리 분석과 전망")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("2026년 5월 기준 · 데이터 범위 2016~2026")
    rs.font.size = Pt(11)
    rs.font.color.rgb = GRAY

    warn = doc.add_paragraph()
    rw = warn.add_run("⚠️ 수치 검증 주의: 2026년 지수 레벨(코스피 8,000대)·투자주체별 정밀 연간 순매수치는 "
                      "복수 언론으로 교차검증했으나 상승 속도가 이례적입니다. 매매 판단 전 KRX 공식데이터"
                      "(data.krx.co.kr MDCSTAT)로 재확인을 권장합니다.")
    rw.italic = True
    rw.font.size = Pt(8.5)
    rw.font.color.rgb = RGBColor(0xB0, 0x30, 0x30)

    # 1. 요약
    heading(doc, "1. 요약 (Executive Summary)")
    bullet(doc, "2025년 코스피 +76%(세계 1위) vs 코스닥 +37% — 약 39%p 역사적 디커플링.", "괴리: ")
    bullet(doc, "삼성전자·SK하이닉스(지수 50%+) HBM/AI 슈퍼사이클이 견인, 코스닥(바이오·2차전지)은 소외.", "원인: ")
    bullet(doc, "코스피=외인·기관 주도 / 코스닥=개인 ~85% 리테일 주도라는 구조적 차이.", "수급: ")
    bullet(doc, "괴리는 구조적이라 단기 완전해소는 어려우나, 국민성장펀드·연기금·소형주 순환매로 부분 키맞추기 진행.", "전망: ")

    # 2. 최근 1년
    heading(doc, "2. 최근 1년 괴리 (2025.6 ~ 2026.5)")
    heading(doc, "2.1 지수 경로", size=11)
    add_table(doc, ["시점", "코스피", "코스피 비고", "코스닥", "코스닥 비고", "비고"],
              INDEX_PATH, widths=[0.9, 0.9, 1.5, 0.9, 1.3, 1.2])
    heading(doc, "2.2 괴리의 정량화", size=11)
    add_table(doc, ["구분", "코스피", "코스닥", "격차/비고"], DIVERGENCE,
              widths=[1.4, 2.1, 1.9, 1.4])

    # 3. 10년 수급
    heading(doc, "3. 10년 투자주체별 수급 (2016~2026)")
    p = doc.add_paragraph()
    p.add_run("시장구조: ").bold = True
    p.add_run("코스피 개인 매매비중 ~57% vs 코스닥 ~85%. 코스닥 개인 지분율 65.1%(소유자의 98.9%가 개인, ~611만명). "
              "→ 코스피=외인·기관, 코스닥=리테일 주도.").font.size = Pt(10)
    heading(doc, "3.1 코스피 외국인 지분율 추이", size=11)
    add_table(doc, ["시점", "외국인 지분율", "비고"], FOREIGN_OWN, widths=[1.5, 1.6, 3.5])
    heading(doc, "3.2 투자주체별 순매수 흐름 (검증치)", size=11)
    add_table(doc, ["시기", "개인", "외국인", "기관", "비고"], FLOWS,
              widths=[1.1, 1.4, 1.6, 1.0, 1.7])
    p2 = doc.add_paragraph()
    p2.add_run("순 내러티브: ").bold = True
    p2.add_run("외국인은 2020~22 구조적 순매도(개인 흡수, 지분 38→30%) → 2024 상반기 밸류업 기대 역대급 매수 복귀 "
               "→ 2025~26 다시 대규모 순매도(2026 YTD 90조+). 단 이번엔 삼성·SK하이닉스 집중 차익실현이라 "
               "지분율은 역설적으로 사상 최고(~39%).").font.size = Pt(9.5)

    # 4. 구조적 요인
    heading(doc, "4. 구조적 요인 (왜 괴리가 구조적인가)")
    add_table(doc, ["요인", "내용", "영향 방향"], DRIVERS, widths=[1.6, 3.4, 1.6])

    # 5. 전망
    heading(doc, "5. 전망")
    heading(doc, "5.1 2026 브로커리지 목표", size=11)
    add_table(doc, ["지수", "기관", "목표", "근거"], FORECAST, widths=[0.8, 1.5, 2.2, 2.1])
    heading(doc, "5.2 종합 전망 (본 리포트 견해)", size=11)
    bullet(doc, "대형주(반도체·밸류업) 우위 유지 + 변동성 확대. 외인 코스피 차익실현↔코스닥 순매수 전환으로 "
                "'괴리 1차 정점 통과' 가능성. 코스닥은 지수보다 중소형 선별 순환매가 현실적.", "단기(2026): ")
    bullet(doc, "완전 해소 조건 = ①코스닥 이익개선·부실정리 ②반도체 외 코스닥 주도주(AI SW·로봇·바이오신약) 출현 "
                "③외인·기관 구조적 복귀. 셋 동반 시 '구조적 괴리→순환적 격차'로 완화.", "중기: ")
    bullet(doc, "삼성·SK하이닉스 HBM 이익방향 / 국민성장펀드·연기금 코스닥 집행 / MSCI 로드맵(2026.6) / "
                "BOK 금리 / 외인 수급교차 지속 여부.", "관전포인트: ")

    # 출처
    heading(doc, "6. 출처")
    add_table(doc, ["내용", "매체", "날짜"], SOURCES, widths=[3.0, 2.4, 1.4])

    path = os.path.join(OUT_DIR, "코스피_코스닥_괴리분석_2026.docx")
    doc.save(path)
    return path


# ----------------------------------------------------------------------------
# EXCEL
# ----------------------------------------------------------------------------
HEAD_FILL = PatternFill("solid", fgColor="1F3564")
HEAD_FONT = Font(bold=True, color="FFFFFF", size=10)
TITLE_FONT = Font(bold=True, color="1F3564", size=13)
THIN = Side(style="thin", color="CCCCCC")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
WRAP = Alignment(wrap_text=True, vertical="top")
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)


def write_sheet(ws, title, headers, rows, widths):
    ws["A1"] = title
    ws["A1"].font = TITLE_FONT
    ws.append([])
    hr = ws.max_row + 1
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=hr, column=c, value=h)
        cell.fill = HEAD_FILL
        cell.font = HEAD_FONT
        cell.alignment = CENTER
        cell.border = BORDER
    for row in rows:
        ws.append(row)
        r = ws.max_row
        for c in range(1, len(headers) + 1):
            cell = ws.cell(row=r, column=c)
            cell.alignment = WRAP
            cell.border = BORDER
            cell.font = Font(size=9.5)
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = ws.cell(row=hr + 1, column=1)


def build_excel():
    wb = Workbook()

    ws = wb.active
    ws.title = "지수경로"
    write_sheet(ws, "코스피·코스닥 지수 경로 (2025.6~2026.5)",
                ["시점", "코스피", "코스피 비고", "코스닥", "코스닥 비고", "비고"],
                INDEX_PATH, [13, 12, 26, 12, 22, 20])

    write_sheet(wb.create_sheet("괴리정량화"), "괴리의 정량화",
                ["구분", "코스피", "코스닥", "격차/비고"], DIVERGENCE, [18, 30, 26, 18])

    write_sheet(wb.create_sheet("외인지분율"), "코스피 외국인 지분율 추이",
                ["시점", "외국인 지분율", "비고"], FOREIGN_OWN, [16, 16, 40])

    write_sheet(wb.create_sheet("투자주체수급"), "투자주체별 순매수 흐름 (2016~2026)",
                ["시기", "개인", "외국인", "기관", "비고"], FLOWS, [14, 22, 28, 16, 32])

    write_sheet(wb.create_sheet("구조적요인"), "구조적 요인",
                ["요인", "내용", "영향 방향"], DRIVERS, [20, 55, 22])

    write_sheet(wb.create_sheet("전망"), "2026 브로커리지 목표 및 전망",
                ["지수", "기관", "목표", "근거"], FORECAST, [10, 18, 28, 30])

    write_sheet(wb.create_sheet("출처"), "출처",
                ["내용", "매체", "날짜"], SOURCES, [38, 34, 14])

    path = os.path.join(OUT_DIR, "코스피_코스닥_데이터_2026.xlsx")
    wb.save(path)
    return path


if __name__ == "__main__":
    w = build_word()
    x = build_excel()
    print("WORD :", w)
    print("EXCEL:", x)
