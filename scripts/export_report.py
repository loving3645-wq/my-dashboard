#!/usr/bin/env python3
"""One-page Word report per ticker.

Reads only from already-collected data/*.json (no live API calls,
zero token cost) and writes reports/{ticker}_{YYYYMMDD}.docx.

Sections:
  - Header: ticker · name · market · sector · timestamp
  - 시세 KPI: 현재가 · 1D / 1W / 1M / 3M / 6M / 1Y 수익률
  - 재무 요약: 3년 매출/매출원가/영업이익/순이익/마진
  - 메자닌: 발행이력 + 본문 파싱 잔액 (있으면)
  - 지분구조: 최대주주 블록 + 소액주주 + Top 5
  - 최근 공시: 30일 / 10건

Usage:
  python scripts/export_report.py 005930
  python scripts/export_report.py 005930 --out reports/samsung.docx
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime, timedelta, timezone

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

KST = timezone(timedelta(hours=9))

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DATA_DIR = os.path.join(ROOT, "data")
REPORTS_DIR = os.path.join(ROOT, "reports")
TICKERS_FILE = os.path.join(ROOT, "tickers-full.js")

KOREAN_FONT = "Malgun Gothic"
ACCENT = RGBColor(0xC8, 0x9B, 0x3C)
INK = RGBColor(0x1A, 0x1A, 0x1A)
INK_MID = RGBColor(0x55, 0x55, 0x55)
INK_FAINT = RGBColor(0x8A, 0x8A, 0x8A)
POS = RGBColor(0x1F, 0x7A, 0x3C)
NEG = RGBColor(0xB0, 0x35, 0x2A)


# ---------- data loaders ----------

def load_json(name):
    path = os.path.join(DATA_DIR, name)
    if not os.path.exists(path):
        return None
    try:
        with open(path, encoding="utf-8") as f:
            return json.load(f)
    except (OSError, json.JSONDecodeError):
        return None


def load_ticker_meta(ticker):
    """Pull name/mkt/sector from tickers-full.js for the requested code."""
    if not os.path.exists(TICKERS_FILE):
        return None
    with open(TICKERS_FILE, encoding="utf-8") as f:
        text = f.read()
    pat = re.compile(
        r"'" + re.escape(ticker) + r"':\s*\{\s*"
        r"name:\s*'([^']+)',\s*"
        r"mkt:\s*'([^']+)',\s*"
        r"sector:\s*'([^']+)'"
    )
    m = pat.search(text)
    if not m:
        return None
    return {"name": m.group(1), "mkt": m.group(2), "sector": m.group(3)}


# ---------- formatting helpers ----------

def fmt_won_eok(won):
    """원 → '5,000억' / '1.20조'."""
    if won is None or not isinstance(won, (int, float)) or won <= 0:
        return "—"
    if won >= 1e12:
        return f"{won / 1e12:.2f}조"
    if won >= 1e8:
        return f"{int(round(won / 1e8)):,}억"
    return f"{int(won):,}원"


def fmt_pct(v, decimals=2):
    if v is None or not isinstance(v, (int, float)):
        return "—"
    return f"{v:+.{decimals}f}%"


def fmt_jo(v):
    """조원 단위 그대로(소수점 2자리)."""
    if v is None or not isinstance(v, (int, float)):
        return "—"
    return f"{v:.2f}조"


def fmt_int(v):
    if v is None:
        return "—"
    try:
        return f"{int(v):,}"
    except (TypeError, ValueError):
        return "—"


# ---------- docx helpers ----------

def _apply_korean_font(run):
    run.font.name = KOREAN_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), KOREAN_FONT)
    rfonts.set(qn("w:ascii"), KOREAN_FONT)
    rfonts.set(qn("w:hAnsi"), KOREAN_FONT)


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_run(p, text, *, size=10, bold=False, color=INK):
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    _apply_korean_font(run)
    return run


def add_para(doc, text="", *, size=10, bold=False, color=INK,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, size=size, bold=bold, color=color)
    return p


def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, size=12, bold=True, color=ACCENT)
    # underline accent bar
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "8")
    bot.set(qn("w:color"), "C89B3C")
    pBdr.append(bot)
    pPr.append(pBdr)


def add_table(doc, headers, rows, *, widths_cm=None, header_fill="F4ECDA"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.autofit = False
    if widths_cm:
        for i, w in enumerate(widths_cm):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, size=9, bold=True, color=INK)
        _set_cell_shading(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        tr = tbl.add_row().cells
        for i, val in enumerate(row):
            text, color, align = _normalize_cell(val)
            tr[i].text = ""
            p = tr[i].paragraphs[0]
            p.alignment = align
            add_run(p, text, size=9, color=color)
            tr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return tbl


def _normalize_cell(val):
    """Cell value can be a str or a (text, color_hint, align_hint) tuple.
    color_hint: 'pos'/'neg'/None.  align_hint: 'left'/'center'/'right'/None.
    """
    if isinstance(val, tuple):
        text, color_hint, align_hint = (val + (None, None))[:3]
    else:
        text, color_hint, align_hint = val, None, None
    color = INK
    if color_hint == "pos":
        color = POS
    elif color_hint == "neg":
        color = NEG
    elif color_hint == "faint":
        color = INK_FAINT
    align = WD_ALIGN_PARAGRAPH.LEFT
    if align_hint == "right":
        align = WD_ALIGN_PARAGRAPH.RIGHT
    elif align_hint == "center":
        align = WD_ALIGN_PARAGRAPH.CENTER
    return str(text), color, align


# ---------- sections ----------

def section_header(doc, ticker, meta, snapshot_entry, now):
    name = meta["name"] if meta else ticker
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, f"{name}", size=20, bold=True, color=INK)
    add_run(p, f"  ({ticker})", size=12, bold=False, color=INK_MID)

    sub_parts = []
    if meta:
        sub_parts.append(meta["mkt"])
        sub_parts.append(meta["sector"])
    sub_parts.append(now.strftime("%Y-%m-%d %H:%M KST"))
    add_para(
        doc,
        " · ".join(sub_parts),
        size=9, color=INK_FAINT, space_after=8,
    )

    # Price / 1D KPI strip
    if snapshot_entry:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        price = snapshot_entry.get("price")
        r1d = snapshot_entry.get("r1d")
        if price:
            add_run(p, f"현재가 {int(price):,}원", size=14, bold=True)
        if isinstance(r1d, (int, float)):
            color = POS if r1d >= 0 else NEG
            add_run(p, f"   {fmt_pct(r1d, 2)} (1D)", size=11, color=color)


def section_returns(doc, snapshot_entry):
    if not snapshot_entry:
        return
    add_section_header(doc, "수익률")
    keys = [("r1d", "1일"), ("r1w", "1주"), ("r1m", "1개월"),
            ("r3m", "3개월"), ("r6m", "6개월"), ("r1y", "1년")]
    headers = [k[1] for k in keys]
    cells = []
    for k, _ in keys:
        v = snapshot_entry.get(k)
        if isinstance(v, (int, float)):
            cells.append((fmt_pct(v, 2), "pos" if v >= 0 else "neg", "right"))
        else:
            cells.append(("—", None, "right"))
    add_table(doc, headers, [cells], widths_cm=[2.4] * 6)


def section_financials(doc, fin_entry):
    if not fin_entry:
        return
    add_section_header(doc, "재무 요약 (DART · 연결우선)")
    years = fin_entry.get("years") or []
    if not years:
        return
    headers = ["항목"] + [f"FY{y}" for y in years]

    def row(label, key, fmt=fmt_jo):
        vals = fin_entry.get(key) or []
        out = [(label, None, "left")]
        for v in vals:
            out.append((fmt(v), None, "right"))
        # pad if short
        while len(out) < len(headers):
            out.append(("—", None, "right"))
        return out

    rows = [
        row("매출액", "revenue"),
        row("매출원가", "cogs"),
        row("매출총이익", "grossProfit"),
        row("영업이익", "operatingIncome"),
        row("당기순이익", "netIncome"),
        [("매출총이익률", None, "left")] + [
            (fmt_pct(v, 1), None, "right") for v in (fin_entry.get("grossMargin") or [])
        ],
        [("영업이익률", None, "left")] + [
            (fmt_pct(v, 1), None, "right") for v in (fin_entry.get("opMargin") or [])
        ],
        [("순이익률", None, "left")] + [
            (fmt_pct(v, 1), None, "right") for v in (fin_entry.get("netMargin") or [])
        ],
        [("매출성장률", None, "left")] + [
            (fmt_pct(v, 1) if v is not None else "—", None, "right")
            for v in (fin_entry.get("revGrowth") or [])
        ],
    ]
    # Pad margin/growth rows to header length
    for r in rows:
        while len(r) < len(headers):
            r.append(("—", None, "right"))
    widths = [3.0] + [2.4] * len(years)
    add_table(doc, headers, rows, widths_cm=widths)


def section_mezzanine(doc, mezz_entry):
    if not mezz_entry or not (mezz_entry.get("issuances") or []):
        return
    add_section_header(doc, "메자닌 발행이력 · 잔액")
    s = mezz_entry.get("summary") or {}
    outstanding = s.get("outstanding")
    if outstanding is None:
        outstanding = s.get("outstandingByMaturity")
    parsed = s.get("outstandingParsedCount") or 0
    parsed_note = (f"{parsed}건 본문 파싱"
                   if parsed > 0 else "본문 파싱 미적용 — 만기 기준")

    summary = (
        f"발행 {s.get('issuanceCount', 0)}건  ·  "
        f"총 발행 {fmt_won_eok(s.get('totalIssued'))}  ·  "
        f"미상환 잔액 {fmt_won_eok(outstanding)} ({parsed_note})  ·  "
        f"행사·상환 이벤트 {s.get('totalEvents', 0)}건"
    )
    next_put = s.get("nextPut")
    if next_put and next_put.get("date"):
        summary += f"  ·  다음 풋 {next_put['date']}"
    add_para(doc, summary, size=9, color=INK_MID, space_after=6)

    headers = ["공시일", "종류", "발행액", "만기", "전환가", "풋(시작)", "잔액"]
    rows = []
    for x in mezz_entry["issuances"]:
        face = x.get("faceAmount")
        cur = x.get("currentOutstanding")
        if cur is not None:
            cur_text = fmt_won_eok(cur)
            cur_hint = "neg" if cur == 0 else ("pos" if face and cur < face else None)
        else:
            cur_text = "—"
            cur_hint = "faint"
        cv_prc = x.get("conversionPrice")
        rows.append([
            (x.get("rceptDt") or "—", None, "left"),
            (x.get("type") or "—", None, "center"),
            (fmt_won_eok(face), None, "right"),
            (x.get("maturityDate") or "—", None, "left"),
            (f"{cv_prc:,}원" if cv_prc else "—", None, "right"),
            (x.get("putStart") or "—", None, "left"),
            (cur_text, cur_hint, "right"),
        ])
    add_table(
        doc, headers, rows,
        widths_cm=[2.0, 1.0, 1.8, 2.0, 2.0, 2.0, 2.0],
    )


def section_ownership(doc, own_entry):
    if not own_entry:
        return
    add_section_header(doc, "지분 구조 (사업보고서 기준)")
    major = own_entry.get("majorBlock") or {}
    minority = own_entry.get("minorityBlock") or {}
    fy = own_entry.get("fiscalYear")
    summary = (
        f"FY{fy}  ·  최대주주(특수관계인 포함) "
        f"{major.get('totalRatio', 0):.2f}%"
    )
    if minority and minority.get("ratio") is not None:
        summary += f"  ·  소액주주 {minority['ratio']:.2f}%"
        if minority.get("shareholderCount"):
            summary += f" ({fmt_int(minority['shareholderCount'])}명)"
    add_para(doc, summary, size=9, color=INK_MID, space_after=6)

    items = (major.get("items") or [])[:5]
    if items:
        headers = ["성명", "관계", "지분율", "보유주식수"]
        rows = [
            [
                (it.get("name") or "—", None, "left"),
                (it.get("relation") or "—", None, "left"),
                (f"{it['ratio']:.2f}%" if it.get("ratio") is not None else "—",
                 None, "right"),
                (fmt_int(it.get("shares")), None, "right"),
            ]
            for it in items
        ]
        add_table(doc, headers, rows, widths_cm=[4.0, 3.0, 2.0, 4.0])


def section_disclosures(doc, disc_list):
    if not disc_list:
        return
    add_section_header(doc, "최근 공시 (30일)")
    headers = ["일자", "구분", "제목"]
    rows = [
        [
            (d.get("date") or "—", None, "left"),
            (d.get("kind") or "—", None, "center"),
            (d.get("title") or "—", None, "left"),
        ]
        for d in disc_list[:10]
    ]
    add_table(doc, headers, rows, widths_cm=[2.4, 1.6, 11.0])


def section_footer(doc, now):
    add_para(doc, "")
    add_para(
        doc,
        f"Generated by 지니 Research Dashboard · {now.strftime('%Y-%m-%d %H:%M KST')}",
        size=8, color=INK_FAINT, align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_after=0,
    )


# ---------- main ----------

def build_report(ticker, out_path=None):
    snap = load_json("market-snapshot.json") or {}
    fin = load_json("financials.json") or {}
    mezz = load_json("mezzanine.json") or {}
    own = load_json("ownership.json") or {}
    disc = load_json("disclosures.json") or {}

    snapshot_entry = (snap.get("tickers") or {}).get(ticker)
    fin_entry = (fin.get("financials") or {}).get(ticker)
    mezz_entry = (mezz.get("mezzanine") or {}).get(ticker)
    own_entry = (own.get("ownership") or {}).get(ticker)
    disc_list = (disc.get("disclosures") or {}).get(ticker)

    meta = load_ticker_meta(ticker)
    now = datetime.now(KST)

    doc = Document()
    # tighter page margins for one-pager density
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    section_header(doc, ticker, meta, snapshot_entry, now)
    section_returns(doc, snapshot_entry)
    section_financials(doc, fin_entry)
    section_mezzanine(doc, mezz_entry)
    section_ownership(doc, own_entry)
    section_disclosures(doc, disc_list)
    section_footer(doc, now)

    if out_path is None:
        os.makedirs(REPORTS_DIR, exist_ok=True)
        stamp = now.strftime("%Y%m%d")
        out_path = os.path.join(REPORTS_DIR, f"{ticker}_{stamp}.docx")
    else:
        os.makedirs(os.path.dirname(os.path.abspath(out_path)) or ".",
                    exist_ok=True)
    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("ticker", help="6-digit KOSPI/KOSDAQ ticker (e.g. 005930)")
    ap.add_argument("--out", help="Output path (default reports/{ticker}_{date}.docx)")
    args = ap.parse_args()

    if not re.match(r"^[0-9A-Z]{6}$", args.ticker):
        print(f"ERROR: invalid ticker format: {args.ticker}", file=sys.stderr)
        sys.exit(2)

    path = build_report(args.ticker, args.out)
    print(f"Wrote {path}")


if __name__ == "__main__":
    main()
